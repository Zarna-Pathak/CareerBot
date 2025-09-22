import streamlit as st
import io
import os
from docx import Document
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader, TextLoader
# from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_google_genai import ChatGoogleGenerativeAI

# --- Load Environment Variables ---
load_dotenv()

# --- Set up the Streamlit Page Configuration ---
st.set_page_config(
    page_title="Resume & Cover Letter Builder",
    layout="wide",
)

# Initialize session state for page control
if 'page' not in st.session_state:
    st.session_state.page = "home"

# Initialize session state for generated cover letter
if 'cover_letter' not in st.session_state:
    st.session_state.cover_letter = None


# ------------------ Helper function: Create DOCX ------------------
def create_docx(resume_text):
    doc = Document()

    def add_formatted_paragraph(line, style=None):
        """ Parse markdown-style bold (**text**) and italics (*text*) and apply formatting in Word paragraphs. """
        p = doc.add_paragraph(style=style)
        while "**" in line:  # Bold handling
            before, bold, rest = line.partition("**")
            if before:
                p.add_run(before)
            bold_text, _, line = rest.partition("**")
            if bold_text:
                p.add_run(bold_text).bold = True
        while "*" in line:  # Italic handling
            before, italic, rest = line.partition("*")
            if before:
                p.add_run(before)
            italic_text, _, line = rest.partition("*")
            if italic_text:
                p.add_run(italic_text).italic = True
        if line:
            p.add_run(line)

    for line in resume_text.splitlines():
        line = line.strip()
        if not line:
            continue
        elif line.lower().startswith(("summary", "skills", "experience", "projects", "education")):
            doc.add_heading(line.replace("**", ""), level=1)
        elif line.startswith("-") or line.startswith("•"):
            add_formatted_paragraph(line.lstrip("-• "), style="List Bullet")
        else:
            add_formatted_paragraph(line)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ------------------ Cover Letter Generator ------------------
def generate_cover_letter(job_description, cv_file):
    """ Parses the CV and generates a tailored cover letter using LangChain. """
    # 1. Parse the CV file
    cv_text = ""
    file_type = cv_file.name.split('.')[-1]

    if file_type == "pdf":
        try:
            temp_file_path = "temp_cv.pdf"
            with open(temp_file_path, "wb") as f:
                f.write(cv_file.getbuffer())
            loader = PyPDFLoader(temp_file_path)
            pages = loader.load_and_split()
            cv_text = "\n".join([page.page_content for page in pages])
            os.remove(temp_file_path)
        except Exception as e:
            st.error(f"Error reading PDF file: {e}")
            return None
    elif file_type == "txt":
        try:
            cv_text = cv_file.getvalue().decode("utf-8")
        except Exception as e:
            st.error(f"Error reading TXT file: {e}")
            return None
    else:
        st.error("Unsupported file type. Please upload a PDF or TXT file.")
        return None

    if not cv_text:
        st.error("Could not extract text from the provided CV.")
        return None

    # 2. Define the LLM
    llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)

    # 3. Prompt
    template = """
    You are an expert career consultant. Your task is to write a compelling cover letter.
    The user will provide you with a Job Description and the content of their Resume (CV).

    Your cover letter should:
    - Be professional, concise, and enthusiastic.
    - Directly reference the candidate's skills and experiences from the CV that are most relevant to the job description.
    - Explain how the candidate is a strong fit for the role.
    - Start with a proper salutation (e.g., "Dear Hiring Manager,").
    - End with a professional closing,
    - Also give the email and social media links if provided in resume.

    --- Job Description: {job_description}
    --- Candidate's Resume: {cv_text}

    --- Generated Cover Letter:
    """

    prompt = PromptTemplate(
        input_variables=["job_description", "cv_text"],
        template=template,
    )

    chain = LLMChain(llm=llm, prompt=prompt)
    response = chain.invoke({"job_description": job_description, "cv_text": cv_text})
    return response['text']


# ------------------ Resume Refiner ------------------
def generate_refined_resume(job_description, cv_file):
    """ Refines the resume based on job description using LLM. """
    cv_text = ""
    file_type = cv_file.name.split('.')[-1]

    if file_type == "pdf":
        try:
            temp_file_path = "temp_resume.pdf"
            with open(temp_file_path, "wb") as f:
                f.write(cv_file.getbuffer())
            loader = PyPDFLoader(temp_file_path)
            pages = loader.load_and_split()
            cv_text = "\n".join([page.page_content for page in pages])
            os.remove(temp_file_path)
        except Exception as e:
            st.error(f"Error reading PDF file: {e}")
            return None
    elif file_type == "txt":
        try:
            cv_text = cv_file.getvalue().decode("utf-8")
        except Exception as e:
            st.error(f"Error reading TXT file: {e}")
            return None
    else:
        st.error("Unsupported file type. Please upload a PDF or TXT file.")
        return None

    if not cv_text:
        st.error("Could not extract text from the provided Resume.")
        return None

    llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.4)

    template = """
    You are a professional career consultant.
    Refine the candidate's resume based on the Job Description provided.

    - Emphasize relevant skills and achievements.
    - Add missing keywords from the JD (if applicable).
    - Rephrase bullet points to align with industry standards.
    - Remove unrelated or less relevant details.
    - Maintain a professional and ATS-friendly format.

    --- Job Description: {job_description}
    --- Original Resume: {cv_text}

    --- Refined Resume:
    """

    prompt = PromptTemplate(
        input_variables=["job_description", "cv_text"],
        template=template,
    )

    chain = LLMChain(llm=llm, prompt=prompt)
    response = chain.invoke({"job_description": job_description, "cv_text": cv_text})
    return response['text']


# ------------------ Resume Ranking ------------------
def rank_resume(job_description, resume_text):
    """ Compare Resume with Job Description and rank it. Show plus and minus points. """
    llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)

    template = """
    You are an expert HR resume reviewer.

    Task: Evaluate the candidate's Resume against the given Job Description.

    Provide:
    - Overall Resume Match Score (out of 100)
    - Plus Points (✅ strengths where the resume matches JD well)
    - Minus Points (❌ missing or weak areas compared to JD)

    --- Job Description: {job_description}
    --- Candidate Resume: {resume_text}

    --- Output in this format:
    Score: XX/100

    ✅ Plus Points:
    - ...
    - ...

    ❌ Minus Points:
    - ...
    - ...
    """

    prompt = PromptTemplate(
        input_variables=["job_description", "resume_text"],
        template=template,
    )

    chain = LLMChain(llm=llm, prompt=prompt)
    response = chain.invoke({"job_description": job_description, "resume_text": resume_text})
    return response['text']


# ------------------ Sidebar ------------------
with st.sidebar:
    st.header("MyApp")
    if st.button("Cover Letter"):
        st.session_state.page = "cover_letter"
    if st.button("Refined Resume"):
        st.session_state.page = "refined_resume"
    if st.button("Rank Resume"):   # <-- new page
        st.session_state.page = "rank_resume"

    st.markdown("---")
    st.markdown("This section will be used to switch between different modules.")


# ------------------ Main Content ------------------
if st.session_state.page == "home":
    st.title("Build Your Resume & Cover Letter 📄")
    st.markdown("Select a module from the sidebar to begin.")

elif st.session_state.page == "cover_letter":
    st.title("Cover Letter Generator")
    st.markdown("Use this tool to create a personalized cover letter tailored to a specific job.")

    if not os.getenv("GOOGLE_API_KEY"):
        st.error("GOOGLE API KEY not found. Please add your key to the .env file.")
    else:
        if st.session_state.cover_letter:
            st.header("Generated Cover Letter Preview")
            st.text_area(label="", value=st.session_state.cover_letter, height=500)

            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("Generate a New Cover Letter"):
                    st.session_state.cover_letter = None
                    st.rerun()
            with col2:
                st.download_button(
                    label="Download Cover Letter",
                    data=st.session_state.cover_letter,
                    file_name="cover_letter.txt",
                    mime="text/plain",
                )
        else:
            job_description = st.text_area(
                "1. Paste the Job Description:",
                height=300,
                placeholder="Paste the full job description here...",
            )
            cv_file = st.file_uploader(
                "2. Upload Your Resume (PDF or TXT):",
                type=["pdf", "txt"],
            )
            if st.button("Generate Cover Letter"):
                if job_description and cv_file:
                    with st.spinner("Generating your cover letter..."):
                        generated_letter = generate_cover_letter(job_description, cv_file)
                        st.session_state.cover_letter = generated_letter
                        st.rerun()
                else:
                    st.error("Please provide both the Job Description and your Resume.")

elif st.session_state.page == "refined_resume":
    st.title("Refined Resume Builder")
    st.markdown("Tailor your resume based on the job description.")

    if not os.getenv("GOOGLE_API_KEY"):
        st.error("GOOGLE API KEY not found.")
    else:
        if 'refined_resume' not in st.session_state:
            st.session_state.refined_resume = None

        if st.session_state.refined_resume:
            st.header("Refined Resume Preview")
            st.text_area("", st.session_state.refined_resume, height=600)

            col1, col2, col3 = st.columns([1, 1, 1])
            with col1:
                if st.button("Generate New Resume"):
                    st.session_state.refined_resume = None
                    st.rerun()
            with col2:
                docx_file = create_docx(st.session_state.refined_resume)
                st.download_button(
                    label="Download Refined Resume (DOCX)",
                    data=docx_file,
                    file_name="refined_resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        else:
            job_description = st.text_area("1. Paste the Job Description:", height=300)
            cv_file = st.file_uploader("2. Upload Your Resume (PDF or TXT):", type=["pdf", "txt"])

            if st.button("Generate Refined Resume"):
                if job_description and cv_file:
                    with st.spinner("Refining your resume..."):
                        refined_resume = generate_refined_resume(job_description, cv_file)
                        st.session_state.refined_resume = refined_resume
                        st.session_state.job_description = job_description
                        st.rerun()
                else:
                    st.error("Please provide both the Job Description and your Resume.")

elif st.session_state.page == "rank_resume":
    st.title("Resume Ranking System")
    st.markdown("Upload your resume and paste the job description to see how well they match.")

    if not os.getenv("GOOGLE_API_KEY"):
        st.error("GOOGLE API KEY not found.")
    else:
        job_description = st.text_area("1. Paste the Job Description:", height=300)
        cv_file = st.file_uploader("2. Upload Your Resume (PDF or TXT):", type=["pdf", "txt"])

        if st.button("Rank Resume"):
            if job_description and cv_file:
                # Extract resume text
                cv_text = ""
                file_type = cv_file.name.split('.')[-1]

                if file_type == "pdf":
                    try:
                        temp_file_path = "temp_rank_resume.pdf"
                        with open(temp_file_path, "wb") as f:
                            f.write(cv_file.getbuffer())
                        loader = PyPDFLoader(temp_file_path)
                        pages = loader.load_and_split()
                        cv_text = "\n".join([page.page_content for page in pages])
                        os.remove(temp_file_path)
                    except Exception as e:
                        st.error(f"Error reading PDF file: {e}")
                        cv_text = None
                elif file_type == "txt":
                    try:
                        cv_text = cv_file.getvalue().decode("utf-8")
                    except Exception as e:
                        st.error(f"Error reading TXT file: {e}")
                        cv_text = None
                else:
                    st.error("Unsupported file type.")
                    cv_text = None

                if cv_text:
                    with st.spinner("Ranking your resume..."):
                        ranking_result = rank_resume(job_description, cv_text)
                    st.subheader("Resume Ranking Result")
                    st.text_area("", ranking_result, height=400)
            else:
                st.error("Please provide both the Job Description and Resume.")

